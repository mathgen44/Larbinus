"""Extraction du texte des documents.

Chaque format a sa fonction, et toutes renvoient la même chose : une liste de
`Fragment`, c'est-à-dire du texte accompagné de sa provenance dans le document
(numéro de page, titre de section, nom de feuille). Cette provenance sert
ensuite à citer précisément la source d'une réponse — un extrait sans origine
n'est pas vérifiable, donc sans valeur.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger("larbinus.rag.extraction")

#: Extensions reconnues, par famille.
EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".xlsm": "xlsx",
    ".html": "html",
    ".htm": "html",
    # Texte en prose : les titres Markdown structurent le découpage.
    ".md": "texte",
    ".markdown": "texte",
    ".txt": "texte",
    ".text": "texte",
    ".csv": "texte",
    ".log": "texte",
    # Code et configuration : traités à part. Les appliquer au lecteur
    # Markdown serait un piège — un `# commentaire` Python ou YAML y passerait
    # pour un titre de section, et l'indentation serait écrasée.
    ".py": "code",
    ".js": "code",
    ".mjs": "code",
    ".ts": "code",
    ".jsx": "code",
    ".tsx": "code",
    ".sh": "code",
    ".bash": "code",
    ".ps1": "code",
    ".sql": "code",
    ".yml": "code",
    ".yaml": "code",
    ".json": "code",
    ".toml": "code",
    ".ini": "code",
    ".cfg": "code",
    ".conf": "code",
    ".go": "code",
    ".rs": "code",
    ".rb": "code",
    ".php": "code",
    ".java": "code",
    ".c": "code",
    ".h": "code",
    ".cpp": "code",
    ".hpp": "code",
    ".css": "code",
}

#: Fichiers reconnus par leur nom, faute d'extension.
NOMS_CONNUS = {
    "dockerfile": "code",
    "makefile": "code",
    "docker-compose.yml": "code",
    "readme": "texte",
}


class FormatNonSupporte(Exception):
    pass


class ExtractionImpossible(Exception):
    pass


@dataclass
class Fragment:
    """Un morceau de texte et son origine dans le document."""

    texte: str
    page: int | None = None
    titre: str | None = None
    meta: dict = field(default_factory=dict)


def famille(nom: str) -> str:
    chemin = Path(nom)
    extension = chemin.suffix.lower()
    if extension in EXTENSIONS:
        return EXTENSIONS[extension]
    # Dockerfile, Makefile… n'ont pas d'extension.
    if chemin.name.lower() in NOMS_CONNUS:
        return NOMS_CONNUS[chemin.name.lower()]
    raise FormatNonSupporte(
        f"Format non pris en charge : « {extension or nom} ». "
        f"Formats acceptés : {', '.join(sorted(EXTENSIONS))}."
    )


def nettoyer(texte: str) -> str:
    """Normalise les espaces sans écraser la structure en paragraphes."""
    texte = texte.replace("\r\n", "\n").replace("\r", "\n")
    texte = re.sub(r"[ \t ]+", " ", texte)
    texte = re.sub(r"\n{3,}", "\n\n", texte)
    return texte.strip()


# --------------------------------------------------------------------------- #
#  Texte brut et Markdown
# --------------------------------------------------------------------------- #
def extraire_texte(donnees: bytes) -> list[Fragment]:
    """Texte brut ou Markdown : les titres `#` découpent le document.

    Conserver le titre de section comme provenance vaut mieux qu'un numéro de
    fragment : c'est ce que l'utilisateur reconnaîtra dans la citation.
    """
    contenu = donnees.decode("utf-8", errors="replace")
    fragments: list[Fragment] = []
    titre_courant: str | None = None
    tampon: list[str] = []

    def vider() -> None:
        texte = nettoyer("\n".join(tampon))
        if texte:
            fragments.append(Fragment(texte=texte, titre=titre_courant))
        tampon.clear()

    for ligne in contenu.split("\n"):
        entete = re.match(r"^(#{1,6})\s+(.*)$", ligne)
        if entete:
            vider()
            titre_courant = entete.group(2).strip()
            continue
        tampon.append(ligne)
    vider()

    if not fragments:
        propre = nettoyer(contenu)
        return [Fragment(texte=propre)] if propre else []
    return fragments


def extraire_code(donnees: bytes) -> list[Fragment]:
    """Code source et fichiers de configuration.

    Aucune détection de titres, et surtout **pas de normalisation des
    espaces** : l'indentation porte du sens en Python comme en YAML, et un
    extrait désindenté serait inutilisable dans une réponse.
    """
    contenu = donnees.decode("utf-8", errors="replace")
    contenu = contenu.replace("\r\n", "\n").replace("\r", "\n")
    lignes = [ligne.rstrip() for ligne in contenu.split("\n")]
    texte = "\n".join(lignes).strip("\n")
    if not texte.strip():
        return []
    return [Fragment(texte=texte, meta={"code": True})]


# --------------------------------------------------------------------------- #
#  PDF
# --------------------------------------------------------------------------- #
def extraire_pdf(donnees: bytes) -> list[Fragment]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dépendance déclarée
        raise ExtractionImpossible("pypdf n'est pas installé.") from exc

    try:
        lecteur = PdfReader(io.BytesIO(donnees))
    except Exception as exc:
        raise ExtractionImpossible(f"PDF illisible : {exc}") from exc

    fragments: list[Fragment] = []
    for numero, page in enumerate(lecteur.pages, start=1):
        try:
            texte = nettoyer(page.extract_text() or "")
        except Exception as exc:
            logger.warning("Page %s illisible : %s", numero, exc)
            continue
        if texte:
            fragments.append(Fragment(texte=texte, page=numero))

    if not fragments:
        # Cas fréquent et déroutant : un PDF de pages scannées ne contient
        # aucun texte. Le dire explicitement évite de chercher ailleurs.
        raise ExtractionImpossible(
            "Aucun texte extractible : ce PDF est probablement un scan "
            "(une reconnaissance de caractères serait nécessaire)."
        )
    return fragments


# --------------------------------------------------------------------------- #
#  Word
# --------------------------------------------------------------------------- #
def extraire_docx(donnees: bytes) -> list[Fragment]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionImpossible("python-docx n'est pas installé.") from exc

    try:
        document = docx.Document(io.BytesIO(donnees))
    except Exception as exc:
        raise ExtractionImpossible(f"Document Word illisible : {exc}") from exc

    fragments: list[Fragment] = []
    titre_courant: str | None = None
    tampon: list[str] = []

    def vider() -> None:
        texte = nettoyer("\n".join(tampon))
        if texte:
            fragments.append(Fragment(texte=texte, titre=titre_courant))
        tampon.clear()

    for paragraphe in document.paragraphs:
        style = (paragraphe.style.name or "").lower()
        if style.startswith(("heading", "titre")) and paragraphe.text.strip():
            vider()
            titre_courant = paragraphe.text.strip()
            continue
        if paragraphe.text.strip():
            tampon.append(paragraphe.text)
    vider()

    # Les tableaux sont aplatis ligne par ligne : la recherche sémantique ne
    # sait rien faire d'une grille, mais une ligne reste une phrase lisible.
    for numero, tableau in enumerate(document.tables, start=1):
        lignes = [
            " | ".join(cellule.text.strip() for cellule in ligne.cells)
            for ligne in tableau.rows
        ]
        texte = nettoyer("\n".join(l for l in lignes if l.strip(" |")))
        if texte:
            fragments.append(Fragment(texte=texte, titre=f"Tableau {numero}"))

    if not fragments:
        raise ExtractionImpossible("Document Word vide ou sans texte.")
    return fragments


# --------------------------------------------------------------------------- #
#  Excel
# --------------------------------------------------------------------------- #
def extraire_xlsx(donnees: bytes, lignes_max: int = 5000) -> list[Fragment]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ExtractionImpossible("openpyxl n'est pas installé.") from exc

    try:
        classeur = load_workbook(io.BytesIO(donnees), read_only=True, data_only=True)
    except Exception as exc:
        raise ExtractionImpossible(f"Classeur illisible : {exc}") from exc

    fragments: list[Fragment] = []
    for feuille in classeur.worksheets:
        lignes: list[str] = []
        for numero, ligne in enumerate(feuille.iter_rows(values_only=True)):
            if numero >= lignes_max:
                logger.warning(
                    "Feuille « %s » tronquée à %s lignes", feuille.title, lignes_max
                )
                break
            cellules = [str(c).strip() for c in ligne if c is not None and str(c).strip()]
            if cellules:
                lignes.append(" | ".join(cellules))
        texte = nettoyer("\n".join(lignes))
        if texte:
            fragments.append(Fragment(texte=texte, titre=feuille.title))
    classeur.close()

    if not fragments:
        raise ExtractionImpossible("Classeur vide.")
    return fragments


# --------------------------------------------------------------------------- #
#  HTML
# --------------------------------------------------------------------------- #
def extraire_html(donnees: bytes) -> list[Fragment]:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:  # pragma: no cover
        raise ExtractionImpossible("beautifulsoup4 n'est pas installé.") from exc

    soupe = BeautifulSoup(donnees.decode("utf-8", errors="replace"), "html.parser")
    for inutile in soupe(["script", "style", "noscript", "nav", "footer", "header"]):
        inutile.decompose()

    titre_page = soupe.title.get_text(strip=True) if soupe.title else None

    fragments: list[Fragment] = []
    titre_courant = titre_page
    tampon: list[str] = []

    def vider() -> None:
        texte = nettoyer("\n".join(tampon))
        if texte:
            fragments.append(Fragment(texte=texte, titre=titre_courant))
        tampon.clear()

    corps = soupe.body or soupe
    for element in corps.find_all(
        ["h1", "h2", "h3", "h4", "p", "li", "pre", "blockquote", "td"]
    ):
        texte = element.get_text(" ", strip=True)
        if not texte:
            continue
        if element.name in ("h1", "h2", "h3", "h4"):
            vider()
            titre_courant = texte
        else:
            tampon.append(texte)
    vider()

    if not fragments:
        texte = nettoyer(soupe.get_text("\n"))
        if not texte:
            raise ExtractionImpossible("Page HTML sans texte exploitable.")
        return [Fragment(texte=texte, titre=titre_page)]
    return fragments


# --------------------------------------------------------------------------- #
#  Aiguillage
# --------------------------------------------------------------------------- #
EXTRACTEURS = {
    "code": extraire_code,
    "pdf": extraire_pdf,
    "docx": extraire_docx,
    "xlsx": extraire_xlsx,
    "html": extraire_html,
    "texte": extraire_texte,
}


def extraire(nom: str, donnees: bytes) -> list[Fragment]:
    """Extrait le texte d'un document, en choisissant le lecteur par extension."""
    return EXTRACTEURS[famille(nom)](donnees)
