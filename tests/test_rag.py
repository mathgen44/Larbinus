"""Socle documentaire du RAG (phase 7a) — aucun appel réseau réel."""

from __future__ import annotations

import io
import zlib

import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.rag.decoupage import contextualiser, decouper
from app.rag.depot import DepotDocuments, cosinus, empreinte
from app.rag.extraction import (
    ExtractionImpossible,
    Fragment,
    FormatNonSupporte,
    extraire,
)
from app.rag.service import ServiceRag
from app.storage.db import Database

pytestmark = pytest.mark.anyio


@pytest.fixture
def anyio_backend():
    return "asyncio"


class EmbeddingsFactices:
    """Vectorise sans réseau : un sac de mots projeté sur 64 dimensions.

    Grossier, mais suffisant pour que deux textes partageant du vocabulaire
    soient plus proches que deux textes étrangers — c'est tout ce que les
    tests ont besoin de vérifier.

    `crc32` et non `hash()` : le hachage des chaînes en Python est randomisé à
    chaque processus, ce qui rendait le classement des résultats variable d'une
    exécution à l'autre. Un test dont le verdict dépend du lancement ne prouve
    rien.
    """

    dimension = 64

    def __init__(self):
        self.appels = 0

    async def vectoriser(self, textes: list[str]) -> list[list[float]]:
        self.appels += 1
        vecteurs = []
        for texte in textes:
            vecteur = [0.0] * self.dimension
            for mot in texte.lower().split():
                vecteur[zlib.crc32(mot.encode()) % self.dimension] += 1.0
            vecteurs.append(vecteur)
        return vecteurs

    async def aclose(self):
        pass


@pytest.fixture
async def rag(tmp_path):
    import app.main as principal

    principal.settings.data_dir = str(tmp_path)
    principal.settings.documents_dir = str(tmp_path / "documents")

    db = Database(tmp_path / "rag.db")
    await db.connect()
    depot = DepotDocuments(db)
    await depot.preparer()
    service = ServiceRag(depot, EmbeddingsFactices(), principal.settings)
    yield service
    await db.close()


# --------------------------------------------------------------------------- #
#  Extraction
# --------------------------------------------------------------------------- #
def test_markdown_decoupe_par_titres():
    contenu = b"""# Installation

Il faut d'abord installer Docker sur la machine cible.

## Configuration reseau

Ouvrir le port 8474 dans le pare-feu.
"""
    fragments = extraire("guide.md", contenu)
    assert [f.titre for f in fragments] == ["Installation", "Configuration reseau"]
    assert "Docker" in fragments[0].texte


def test_html_ignore_scripts_et_navigation():
    html = b"""<html><head><title>Page</title></head><body>
    <nav>Menu inutile</nav>
    <script>var poison = 1;</script>
    <h1>Titre</h1><p>Contenu utile.</p>
    <footer>Pied de page</footer></body></html>"""
    fragments = extraire("page.html", html)
    texte = " ".join(f.texte for f in fragments)
    assert "Contenu utile" in texte
    assert "poison" not in texte
    assert "Menu inutile" not in texte


def test_format_inconnu_refuse():
    with pytest.raises(FormatNonSupporte):
        extraire("archive.zip", b"...")


def test_pdf_scanne_donne_un_message_explicite():
    """Un PDF sans couche texte est le piège le plus courant."""
    from pypdf import PdfWriter

    tampon = io.BytesIO()
    ecrivain = PdfWriter()
    ecrivain.add_blank_page(width=200, height=200)
    ecrivain.write(tampon)

    with pytest.raises(ExtractionImpossible) as exc:
        extraire("scan.pdf", tampon.getvalue())
    assert "scan" in str(exc.value).lower()


def test_docx_et_xlsx():
    import docx
    from openpyxl import Workbook

    document = docx.Document()
    document.add_heading("Procedure", level=1)
    document.add_paragraph("Redemarrer le service avant toute chose.")
    tampon = io.BytesIO()
    document.save(tampon)
    fragments = extraire("procedure.docx", tampon.getvalue())
    assert fragments[0].titre == "Procedure"
    assert "Redemarrer" in fragments[0].texte

    classeur = Workbook()
    feuille = classeur.active
    feuille.title = "Inventaire"
    feuille.append(["Machine", "Role"])
    feuille.append(["beast", "Proxmox"])
    tampon = io.BytesIO()
    classeur.save(tampon)
    fragments = extraire("inventaire.xlsx", tampon.getvalue())
    assert fragments[0].titre == "Inventaire"
    assert "beast" in fragments[0].texte


def test_code_nest_pas_lu_comme_du_markdown():
    """Un `#` en Python est un commentaire, pas un titre de section."""
    source = b'''# Configuration du service
import os

def demarrer(port: int = 8474):
    if port < 1024:
        raise ValueError("port reserve")
    return os.environ.get("HOTE", "0.0.0.0"), port
'''
    fragments = extraire("service.py", source)
    assert len(fragments) == 1
    assert fragments[0].titre is None          # aucun faux titre
    assert fragments[0].meta.get("code") is True
    # L'indentation doit survivre : elle porte la structure du code.
    assert "    if port < 1024:" in fragments[0].texte


def test_yaml_et_fichiers_sans_extension():
    fragments = extraire("docker-compose.yml", b"services:\n  larbinus:\n    build: .\n")
    assert fragments[0].meta.get("code") is True
    assert "  larbinus:" in fragments[0].texte

    fragments = extraire("Dockerfile", b"FROM python:3.12-slim\nWORKDIR /app\n")
    assert fragments[0].meta.get("code") is True


def test_code_decoupe_par_lignes_entieres():
    from app.rag.decoupage import decouper as _decouper

    lignes = "\n".join(f"    ligne_{i} = calcul({i})" for i in range(120))
    morceaux = _decouper([Fragment(texte=lignes, meta={"code": True})], taille=400)
    assert len(morceaux) > 1
    # Aucune ligne ne doit avoir été coupée en son milieu.
    for morceau in morceaux:
        for ligne in morceau.contenu.split("\n"):
            assert ligne == "" or ligne.startswith("    ligne_")


# --------------------------------------------------------------------------- #
#  Découpage
# --------------------------------------------------------------------------- #
def test_decoupage_respecte_les_phrases_et_chevauche():
    phrase = "Ceci est une phrase de test suffisamment longue pour compter. "
    fragments = [Fragment(texte=phrase * 40, titre="Section")]
    morceaux = decouper(fragments, taille=300, chevauchement=100)

    assert len(morceaux) > 1
    assert all(m.titre == "Section" for m in morceaux)
    # Aucun morceau ne doit dépasser franchement la cible.
    assert all(len(m.contenu) <= 420 for m in morceaux)
    # Le chevauchement doit faire réapparaître la fin du morceau précédent.
    assert morceaux[0].contenu.split(".")[-2].strip() in morceaux[1].contenu


def test_phrase_plus_longue_que_la_cible_est_coupee():
    morceaux = decouper([Fragment(texte="a" * 900)], taille=300, chevauchement=50)
    assert len(morceaux) == 3


def test_contexte_porte_la_provenance():
    morceaux = decouper([Fragment(texte="Le port est 8474." * 20, titre="Reseau", page=3)])
    texte = contextualiser(morceaux[0], "guide.pdf")
    assert texte.startswith("[guide.pdf — Reseau — page 3]")


# --------------------------------------------------------------------------- #
#  Dépôt et recherche
# --------------------------------------------------------------------------- #
async def test_indexation_et_recherche(rag):
    await rag.deposer(
        "reseau.md",
        "# Pare-feu\n\nLe port 8474 doit etre ouvert pour Larbinus.\n".encode(),
    )
    await rag.deposer(
        "cuisine.md",
        "# Tarte\n\nMelanger la farine et le beurre puis cuire vingt minutes.\n".encode(),
    )

    documents = await rag.depot.documents()
    assert {d["status"] for d in documents} == {"indexe"}
    assert all(d["chunk_count"] >= 1 for d in documents)

    resultats = await rag.rechercher("Quel port ouvrir dans le pare-feu ?", limite=1)
    assert resultats and resultats[0]["filename"] == "reseau.md"


async def test_meme_fichier_depose_deux_fois_n_est_pas_duplique(rag):
    donnees = b"# Note\n\nContenu quelconque mais assez long pour etre indexe.\n"
    premier = await rag.deposer("note.md", donnees)
    second = await rag.deposer("note-copie.md", donnees)

    assert second["doublon"] is True
    assert second["id"] == premier["id"]
    assert len(await rag.depot.documents()) == 1


async def test_reindexation_apres_echec_du_service(rag):
    """Le cas réel : le modèle d'embedding manquait, on vient de l'installer."""
    rag.client = None
    document = await rag.deposer(
        "note.md", b"# Titre\n\nUn contenu assez long pour produire un fragment.\n"
    )
    assert document["status"] == "en_attente"

    # Le fichier déposé doit avoir été conservé, sinon aucune reprise possible.
    assert (rag.dossier_depots / document["path"]).is_file()

    rag.client = EmbeddingsFactices()
    resume = await rag.reindexer_les_echecs()
    assert resume == {"tentes": 1, "indexes": 1}

    recharge = await rag.depot.document(document["id"])
    assert recharge["status"] == "indexe" and recharge["chunk_count"] >= 1
    assert await rag.rechercher("contenu", limite=1)


async def test_suppression_retire_aussi_le_fichier_conserve(rag):
    document = await rag.deposer(
        "jetable.md", b"# Titre\n\nUn contenu assez long pour produire un fragment.\n"
    )
    chemin = rag.dossier_depots / document["path"]
    assert chemin.is_file()

    assert await rag.supprimer(document["id"]) is True
    assert not chemin.exists()
    assert await rag.depot.documents() == []


async def test_document_illisible_est_marque_en_erreur_sans_lever(rag):
    resultat = await rag.deposer("photo.png", b"\x89PNG\r\n\x1a\n")
    assert resultat["status"] == "erreur"
    assert "non pris en charge" in resultat["error"].lower()


async def test_suppression_retire_fragments_et_vecteurs(rag):
    document = await rag.deposer(
        "a-supprimer.md", b"# Titre\n\nUn contenu assez long pour produire un fragment.\n"
    )
    assert await rag.rechercher("contenu", limite=5)

    await rag.depot.supprimer_document(document["id"])
    assert await rag.depot.documents() == []
    assert await rag.rechercher("contenu", limite=5) == []


async def test_contexte_numerote_et_cite_les_sources(rag):
    await rag.deposer(
        "procedure.md",
        "# Redemarrage\n\nPour redemarrer le conteneur, lancer docker compose restart.\n".encode(),
    )
    contexte, sources = await rag.contexte("Comment redemarrer le conteneur ?", limite=1)

    assert "[1]" in contexte
    assert "docker compose restart" in contexte
    # La consigne anti-invention doit accompagner les extraits.
    assert "dis-le clairement" in contexte
    assert sources[0]["filename"] == "procedure.md"
    assert sources[0]["numero"] == 1


async def test_sans_embeddings_le_document_reste_en_attente(rag):
    rag.client = None
    resultat = await rag.deposer("note.md", b"# Titre\n\nUn contenu quelconque ici.\n")
    assert resultat["status"] == "en_attente"
    assert "EMBEDDING_PROVIDER" in resultat["error"]
    assert await rag.rechercher("contenu") == []


async def test_changement_de_modele_detecte(rag):
    """Des vecteurs de dimensions différentes ne sont pas comparables."""
    from app.rag.depot import IndexIncoherent

    await rag.deposer("note.md", b"# Titre\n\nUn contenu suffisamment long ici.\n")
    with pytest.raises(IndexIncoherent):
        await rag.depot.rechercher([0.0] * 8, limite=3)


async def test_repli_sans_sqlite_vec(tmp_path):
    """Le filet de sécurité doit fonctionner : même résultats, sans l'extension."""
    import app.main as principal

    principal.settings.data_dir = str(tmp_path)
    db = Database(tmp_path / "repli.db")
    await db.connect()
    depot = DepotDocuments(db)
    await depot.preparer()
    # On force le repli en Python, comme si l'extension n'avait pas pu être chargée.
    depot.vec_disponible = False

    service = ServiceRag(depot, EmbeddingsFactices(), principal.settings)
    await service.deposer(
        "reseau.md", b"# Pare-feu\n\nLe port 8474 doit etre ouvert pour Larbinus.\n"
    )
    await service.deposer(
        "cuisine.md", b"# Tarte\n\nMelanger la farine et le beurre puis cuire.\n"
    )

    etat = await depot.etat_index()
    assert etat["moteur"] == "python"

    resultats = await service.rechercher("Quel port ouvrir dans le pare-feu ?", limite=1)
    assert resultats and resultats[0]["filename"] == "reseau.md"
    assert 0.0 <= resultats[0]["score"] <= 1.0
    await db.close()


# --------------------------------------------------------------------------- #
#  Dossier surveillé
# --------------------------------------------------------------------------- #
async def test_scan_du_dossier_surveille(rag, tmp_path):
    dossier = tmp_path / "documents"
    (dossier / "sous-dossier").mkdir(parents=True)
    (dossier / "note.md").write_text("# Note\n\nUn contenu assez long pour l'index.\n")
    (dossier / "sous-dossier" / "autre.txt").write_text(
        "Un second document, dans un sous-dossier, avec du texte."
    )
    (dossier / "image.png").write_bytes(b"\x89PNG")

    resume = await rag.scanner()
    assert resume["existe"] is True
    assert resume["ajoutes"] == 2
    assert resume["ignores"] == 1        # le PNG n'est pas un format reconnu

    # Un second scan ne réindexe pas ce qui n'a pas changé.
    second = await rag.scanner()
    assert second["ajoutes"] == 0 and second["inchanges"] == 2

    documents = await rag.depot.documents()
    assert {d["source"] for d in documents} == {"dossier"}
    assert any(d["path"].endswith("autre.txt") for d in documents)


async def test_dossier_absent_est_signale(rag, tmp_path):
    rag.settings.documents_dir = str(tmp_path / "inexistant")
    resume = await rag.scanner()
    assert resume["existe"] is False
    assert "montez-le" in resume["message"]


# --------------------------------------------------------------------------- #
#  Divers
# --------------------------------------------------------------------------- #
def test_empreinte_et_cosinus():
    assert empreinte(b"abc") == empreinte(b"abc")
    assert empreinte(b"abc") != empreinte(b"abd")
    assert cosinus([1, 0], [1, 0]) == pytest.approx(1.0)
    assert cosinus([1, 0], [0, 1]) == pytest.approx(0.0)
    assert cosinus([0, 0], [1, 0]) == 0.0


# --------------------------------------------------------------------------- #
#  Routes
# --------------------------------------------------------------------------- #
def test_routes_documents(tmp_path):
    import app.main as principal

    principal.settings.data_dir = str(tmp_path)
    principal.settings.documents_dir = str(tmp_path / "documents")

    with TestClient(app) as client:
        app.state.rag.client = EmbeddingsFactices()

        etat = client.get("/api/documents").json()
        assert etat["disponible"] is True
        assert etat["documents"] == []
        assert ".pdf" in etat["formats"]

        depot = client.post(
            "/api/documents",
            files={"files": ("note.md", b"# Titre\n\nUn contenu assez long pour l'index.\n",
                             "text/markdown")},
        )
        assert depot.status_code == 201
        document = depot.json()["resultats"][0]
        assert document["status"] == "indexe"

        recherche = client.get("/api/documents/recherche", params={"q": "contenu"})
        assert recherche.status_code == 200
        assert recherche.json()["resultats"][0]["filename"] == "note.md"

        assert client.delete(f"/api/documents/{document['id']}").status_code == 204
        assert client.get("/api/documents").json()["documents"] == []


# --------------------------------------------------------------------------- #
#  Intégration au chat (phase 7b)
# --------------------------------------------------------------------------- #
@pytest.fixture
def client_rag(tmp_path):
    """Client complet : Ollama simulé, embeddings factices, un document indexé."""
    import httpx

    import app.main as principal
    from app.config import Settings
    from app.providers.registry import ProviderRegistry
    from tests.conftest import ndjson, patch_provider

    principal.settings.data_dir = str(tmp_path)
    principal.settings.documents_dir = str(tmp_path / "documents")

    with TestClient(app) as testeur:
        app.state.rag.client = EmbeddingsFactices()
        registry = ProviderRegistry(
            Settings(_env_file=None, ollama_base_url="http://ollama.test:11434")
        )
        patch_provider(
            registry.get("ollama"),
            {
                "/api/chat": httpx.Response(
                    200,
                    content=ndjson(
                        {"message": {"content": "Le port est 8474 [1]."}, "done": False},
                        {"done": True, "prompt_eval_count": 5, "eval_count": 6},
                    ),
                ),
                "/api/tags": httpx.Response(200, json={"models": [{"name": "mistral"}]}),
            },
        )
        app.state.registry = registry

        testeur.post(
            "/api/documents",
            files={"files": ("reseau.md",
                             b"# Pare-feu\n\nLe port 8474 doit etre ouvert pour Larbinus.\n",
                             "text/markdown")},
        )
        yield testeur


def test_chat_sans_rag_n_injecte_rien(client_rag):
    conversation = client_rag.post("/api/conversations", json={}).json()
    assert conversation["rag"] == 0

    reponse = client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir ?"}],
              "stream": False},
    ).json()
    assert "sources" not in reponse


def test_chat_avec_rag_cite_ses_sources(client_rag):
    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    assert conversation["rag"] == 1

    reponse = client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir dans le pare-feu ?"}],
              "stream": False},
    ).json()

    assert reponse["sources"][0]["filename"] == "reseau.md"
    assert reponse["sources"][0]["numero"] == 1
    assert "8474" in reponse["sources"][0]["extrait"]


def test_les_extraits_arrivent_dans_le_prompt_systeme(client_rag):
    """Vérifie ce que le modèle reçoit réellement, pas seulement la réponse."""
    import json as _json

    import httpx

    from tests.conftest import ndjson

    envoye = {}
    provider = app.state.registry.get("ollama")
    original = provider._client

    def intercepte(request: httpx.Request) -> httpx.Response:
        envoye.update(_json.loads(request.content))
        return httpx.Response(
            200, content=ndjson({"message": {"content": "ok"}, "done": False}, {"done": True})
        )

    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    provider._client = httpx.AsyncClient(transport=httpx.MockTransport(intercepte))
    try:
        client_rag.post(
            "/api/chat",
            json={"model": "ollama/mistral", "conversation_id": conversation["id"],
                  "messages": [{"role": "user", "content": "Quel port ouvrir dans le pare-feu ?"}],
                  "stream": False},
        )
    finally:
        provider._client = original

    systeme = envoye["messages"][0]
    assert systeme["role"] == "system"
    assert "8474" in systeme["content"]
    assert "[1]" in systeme["content"]


def test_evenement_sources_avant_la_reponse(client_rag):
    """L'utilisateur doit voir les sources pendant la rédaction, pas après."""
    import json as _json

    from tests.test_chat import parse_sse

    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    reponse = client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir dans le pare-feu ?"}]},
    )
    evenements = parse_sse(reponse.text)
    noms = [nom for nom, _ in evenements]
    assert noms.index("sources") < noms.index("delta")

    charge = _json.loads(next(d for nom, d in evenements if nom == "sources"))
    assert charge["sources"][0]["filename"] == "reseau.md"


def test_sources_conservees_avec_la_reponse(client_rag):
    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir dans le pare-feu ?"}],
              "stream": False},
    )
    messages = client_rag.get(f"/api/conversations/{conversation['id']}").json()["messages"]
    assert messages[1]["sources"][0]["filename"] == "reseau.md"


def test_le_contexte_ne_pollue_pas_l_historique(client_rag):
    """Les extraits valent pour un tour : ils ne doivent pas s'accumuler."""
    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    for question in ["Quel port ouvrir ?", "Et pour le pare-feu ?"]:
        client_rag.post(
            "/api/chat",
            json={"model": "ollama/mistral", "conversation_id": conversation["id"],
                  "messages": [{"role": "user", "content": question}], "stream": False},
        )
    messages = client_rag.get(f"/api/conversations/{conversation['id']}").json()["messages"]
    assert all("Extraits des documents" not in m["content"] for m in messages)


def test_le_champ_rag_de_la_requete_prime(client_rag):
    """Une conversation sans RAG peut l'activer ponctuellement, et l'inverse."""
    conversation = client_rag.post("/api/conversations", json={}).json()
    reponse = client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir dans le pare-feu ?"}],
              "rag": True, "stream": False},
    ).json()
    assert "sources" in reponse


def test_index_en_panne_n_empeche_pas_de_repondre(client_rag):
    """Une recherche impossible dégrade la réponse, elle ne la bloque pas."""
    from app.rag.embeddings import EmbeddingIndisponible

    class ClientEnPanne:
        async def vectoriser(self, textes):
            raise EmbeddingIndisponible("service arrêté")

        async def aclose(self):
            pass

    app.state.rag.client = ClientEnPanne()
    conversation = client_rag.post("/api/conversations", json={"rag": True}).json()
    reponse = client_rag.post(
        "/api/chat",
        json={"model": "ollama/mistral", "conversation_id": conversation["id"],
              "messages": [{"role": "user", "content": "Quel port ouvrir ?"}],
              "stream": False},
    )
    assert reponse.status_code == 200
    assert "sources" not in reponse.json()
    assert reponse.json()["content"]
